import docx
import os
from django.conf import settings
from django.utils.timezone import now


def generate_document(author: str, institution: str, body: str, date=now()):
    documents_path = os.path.join(settings.BASE_DIR, 'apps', 'revista_cientifica', 'documents')
    base_document = docx.Document(documents_path+f'/BaseDocument.docx')
    base_document.paragraphs[14].text = f'La Habana, {now().year}/ {now().month}/ {now().day}'
    base_document.paragraphs[17].text = f'{author}'
    base_document.paragraphs[18].text = f'{institution}'
    base_document.paragraphs[24].text = f'{body}'
    file_name = f'/GeneratedAt:{now().year}-{now().month}-{now().day} {now().hour}:{now().minute}:{now().second}.docx'
    base_document.save(documents_path + file_name)
    return documents_path + file_name
